from docx import Document
from docx.shared import Inches
from io import BytesIO
from .base import ReportGenerator


class DocxReportGenerator(ReportGenerator):
    def get_headers(self, report_type):
        headers = {
            'appointments': [
                ('appointment_date', 'Fecha'),
                ('patient_name', 'Paciente'),
                ('doctor_name', 'Doctor'),
                ('speciality__name', 'Especialidad'),
                ('status', 'Estado')
            ],
            'patients': [
                ('full_name', 'Nombre'),
                ('email', 'Correo'),
                ('identification', 'Cédula'),
                ('phone_number', 'Teléfono'),
                ('birth_date', 'Fecha Nacimiento'),
                ('genre', 'Género')
            ],
            'doctors': [
                ('full_name', 'Nombre Completo'),
                ('email', 'Correo'),
                ('identification', 'Cédula'),
                ('phone_number', 'Teléfono'),
                ('speciality__name', 'Especialidades')
            ]
        }
        return headers.get(report_type, [])

    def format_data(self, item, report_type):
        if report_type == 'appointments':
            return {
                'appointment_date': f"{item['appointment_date'].strftime('%d/%m/%Y')} {item['appointment_time'].strftime('%H:%M')}",
                'patient_name': f"{item['patient__first_name']} {item['patient__last_name']}",
                'doctor_name': f"{item['doctor__first_name']} {item['doctor__last_name']}",
                'speciality__name': item['speciality__name'],
                'status': self.translate_status(item['status'])
            }
        elif report_type == 'patients':
            return {
                'full_name': f"{item['first_name']} {item['last_name']}",
                'email': item['email'],
                'identification': item['identification'],
                'phone_number': item['phone_number'],
                'birth_date': item['birth_date'].strftime('%d/%m/%Y') if item['birth_date'] else '',
                'genre': 'Femenino' if item['genre'] == 'F' else 'Masculino'
            }
        elif report_type == 'doctors':
            return {
                'full_name': f"{item['first_name']} {item['last_name']}",
                'email': item['email'],
                'identification': item['identification'],
                'phone_number': item['phone_number'],
                'speciality__name': item['specialities__name']
            }
        return item

    def translate_status(self, status):
        status_map = {
            'PENDING': 'Pendiente',
            'CONFIRMED': 'Confirmada',
            'CANCELLED': 'Cancelada',
            'ATTENDED': 'Atendida',
            'NO_SHOW': 'No asistió'
        }
        return status_map.get(status, status)

    def get_report_title(self, report_type):
        titles = {
            'appointments': 'Reporte de Citas',
            'patients': 'Reporte de Pacientes',
            'doctors': 'Reporte de Médicos'
        }
        return titles.get(report_type, f'Reporte de {report_type}')

    def generate(self, data, template=None):
        doc = Document()
        report_type = template.split('/')[-1].split('.')[0]
        headers = self.get_headers(report_type)

        # Add title
        doc.add_heading(self.get_report_title(report_type), 0)

        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Add headers
        for idx, (_, header) in enumerate(headers):
            table.cell(0, idx).text = header

        # Add data rows
        for item in data:
            formatted_item = self.format_data(item, report_type)
            row_cells = table.add_row().cells
            for idx, (key, _) in enumerate(headers):
                value = formatted_item.get(key, '')
                row_cells[idx].text = str(value)

        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.getvalue()
